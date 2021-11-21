from docx import Document
from docx.shared import Inches


class invoiceGenerator:
    def __init__(self,personName,projectName,price):
        self.personName=personName
        self.projectName=projectName
        self.price=price
    
    def generate(self):
        invoice = Document()
        file1=open("invoiceNumber.txt",'r+')
        invoiceNumber=int(file1.read())
        file1.seek(0)
        file1.write(str(invoiceNumber+1))
        file1.close()

        invoice.add_picture('LookOut.png',width=Inches(4),height=Inches (3))
        invoice.add_heading("Invoice: {}".format(self.projectName),level=0 )
        invoice.add_heading("Invoice Number: {}-{}".format(invoiceNumber,self.personName[:3]),level=1 )
        invoice.add_heading("Client: {}".format(self.personName),level=1 )
        invoice.add_heading("",level=0 )
        invoice.add_heading("Payment Due: ${:.2f}".format(self.price),level=0 )


        invoice.save('Invoice-{}_{}-{}.docx'.format(self.projectName,invoiceNumber,self.personName[:3]))


        
